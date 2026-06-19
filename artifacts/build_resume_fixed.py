from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import HRFlowable, KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(r"E:\zhishu\zhishu1\artifacts")
DOCX_PATH = ROOT / "resume_dengke_optimized.docx"
PDF_PATH = ROOT / "resume_dengke_optimized.pdf"

ACCENT = RGBColor(28, 73, 123)
TEXT = RGBColor(26, 35, 48)
MUTED = RGBColor(90, 102, 118)

NAME = "邓柯"
META = "深圳技术大学｜数据科学与大数据技术｜2024级本科在读｜电话：19966348826｜邮箱：dengke051105@gmail.com｜地点：广东 深圳"
INTRO = (
    "具备扎实的数理逻辑推演基础与独立工程实践经验，熟练掌握 Python、JavaScript、TypeScript、C/C++。"
    "能够覆盖从数据获取、训练数据构造、轻量模型验证，到前后端系统交互、数据库建模、接口联调与可视化展示的完整开发链路，"
    "擅长将底层算法逻辑转化为可交付的工程化系统。"
)
EDU = "2024.09 - 至今｜深圳技术大学｜数据科学与大数据技术本科"
EDU_DETAIL = (
    "平均绩点 3.58/4.50，专业排名 17/150。主修数据结构、数据科学基础、人工智能数学基础、"
    "大数据原理与技术、操作系统、计算机组成与系统结构、数据库系统、机器学习导论等。"
)

INFO_ROWS = [
    [("出生年月", "2005/11/05"), ("政治面貌", "共青团员"), ("在读院校", "深圳技术大学")],
    [("专业", "数据科学与大数据技术"), ("入学年份", "2024"), ("毕业年份", "2028")],
]

SKILLS = [
    "编程语言：Python、JavaScript、TypeScript、C/C++",
    "工程与框架：Playwright、Next.js、React、Tailwind CSS、Shadcn UI、Supabase、Vercel",
    "数据与模型：数据清洗 Pipeline、时序建模、合成数据构造、条件概率逻辑、模型指标分析",
    "前端可视化：TypeScript 大屏开发、ECharts、Three.js、GLTFLoader、实时流式数据展示",
]

PROJECTS = [
    (
        "实践一：知枢——企业知识资产管理与智能检索平台",
        "围绕数据库系统课程答辩目标，参与并推进知识库系统从本地演示到 MySQL 主链路可运行 Demo 的落地，"
        "串联文档、版本、FAQ、用户、阅读收藏、问答引用与 Agent 留痕等核心数据实体；完成前后端接口联调与页面交互实现，"
        "补齐登录鉴权、刷新令牌、角色权限、文档分段重建、文件上传下载与对象存储镜像等功能；编写 MySQL、路线 B 与 Qdrant Demo "
        "验收脚本，打通本地数据库、向量检索演示与课程展示页面，支撑项目完成整链路验证与答辩演示。"
    ),
    (
        "实践二：“猎穹科技”电商小程序开发",
        "负责微信电商小程序前端页面与核心交互链路开发，围绕首页、商品详情、购物车、结算、订单、个人中心及管理端页面完成界面实现与业务状态联动；"
        "主导前后端接口联调与数据适配，处理登录鉴权、Token 失效兜底、商品与购物车状态同步、订单确认收货回调等业务流程；"
        "针对首页商品图预热、登录等待时长、费用明细展示等体验问题进行优化，提升页面加载效率与下单链路完整性。"
    ),
    (
        "实践三：问卷自动化获取与数据分析模型",
        "基于 Playwright 构建自动化控制脚本，实现复杂表单 DOM 节点的动态解析与精准交互；设计基于历史样本的动态权重分配算法，"
        "结合条件概率逻辑与带权随机采样，模拟高真实度数据生成行为；构建数据清洗与特征提取管道，输出规范化数据集。"
    ),
    (
        "实践四：“our-space”全栈交互网站开发",
        "采用 Vercel 完成仓库托管、自动化构建与持续部署，接入 Supabase 作为后端服务支撑；独立设计并封装数据抽取、用户留言等双向交互接口，"
        "处理前后端数据通信与状态同步逻辑。"
    ),
    (
        "实践五：食安大模型 Agent 集群——中枢路由 Agent",
        "负责食安大模型 Agent 集群的用户交互入口与中枢路由层设计实现，主导定义路由层全局 Schema、前后端 HTTP / WebSocket 协议及下游 Agent 接口边界；"
        "搭建 Router 到 DataAgent / TriageAgent 的事件驱动分发链路，完善会话状态、确认流与 pending 到 final 的结果回传机制；"
        "编写中枢 Agent 的 System Prompt、Function Calling 工具与路由策略，并推动前端从联调页迭代为面向用户的极简聊天页；"
        "围绕多轮查询上下文继承问题，沉淀结构化状态 carry-over、query rewriting 与下游约束下传方案。"
    ),
    (
        "实践六：计算机设计大赛——Aura-Audit（代码安全与免疫平台）",
        "全栈独立开发与架构决策，采用 Next.js、React、Tailwind CSS 与 Shadcn UI 进行工程开发，引入神经符号（Neuro-symbolic）AI 技术作为核心审计引擎。"
    ),
    (
        "实践七：国创赛——天枢：电驱系统健康管理平台",
        "负责平台前端主视图设计、前后端联调与预测展示闭环搭建；使用 TypeScript 构建总览大屏、传感器详情页与预测详情页，完成实时数据快照、"
        "SSE 流式更新、趋势分析与风险预警展示；基于合成故障时序数据补充 Transformer 风格轻量预测方案，参与训练数据方案梳理与模型结果接入；"
        "使用 Three.js 与 GLTFLoader 落地直升机电驱系统三维展示，支持真实外观 / 剖透双模式、部件聚焦放大与归位；同步沉淀前后端接口对接文档与训练方案文档，"
        "提升项目可复现性与大会展示完整度。"
    ),
    (
        "实践八：doScenes Challenge——语言条件自动驾驶轨迹预测",
        "负责 doScenes Challenge 语言 + 历史轨迹预测方案的训练、评估与提交工程化；基于 PyTorch 与 Transformers 搭建正式训练、语言增益诊断、"
        "官方 submission / precheck 与技术报告生成链路，统一 ADE / FDE 的 meters 口径并对齐官方 127 行 scene_token 提交协议；"
        "围绕 cross-attention、future-query residual 解码、warm-start 迁移等方向持续迭代模型结构与训练配置，完成多轮 ADE 误差诊断与方案优化，"
        "实际比赛提交达到 ADE 3.2107、FDE 7.2955。"
    ),
    (
        "实践九：东方风格弹幕射击游戏 UI 状态优化（C++ / SDL2）",
        "在 C++ 面向对象课程项目中负责游戏界面流程与状态管理模块优化，基于 SDL2 对主状态机进行整理，新增帮助界面与暂停界面，"
        "完善主菜单、角色选择、战斗、对话、暂停及结算等场景间的切换逻辑；实现主菜单帮助入口、战斗中 ESC 暂停、暂停后继续游戏或返回主菜单等关键交互，"
        "并统一相关输入响应与界面渲染逻辑，提升了项目演示完整度、代码可维护性与课程答辩展示效果。"
    ),
]

COMPETITIONS = [
    "2025 全国大学生数学建模广东赛区二等奖（负责编程工作）",
    "2025 全国大学生数学竞赛广东赛区三等奖",
    "2025 挑战杯校级优胜奖（垃圾分类识别，训练 YOLOv7 模型）",
]


def set_run_font(run, name="Microsoft YaHei", size=10.5, bold=False, color=TEXT):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def format_paragraph(paragraph, space_before=0, space_after=0, line=1.12, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = line
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if align is not None:
        paragraph.alignment = align


def add_section_border(paragraph):
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B7CBE3")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = TEXT

    for style_name in ["ResumeSection", "ProjectTitle", "MetaLine"]:
        if style_name not in styles:
            styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    section_style = styles["ResumeSection"]
    section_style.base_style = normal
    section_style.font.name = "Microsoft YaHei"
    section_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    section_style.font.size = Pt(12.2)
    section_style.font.bold = True
    section_style.font.color.rgb = ACCENT

    project_style = styles["ProjectTitle"]
    project_style.base_style = normal
    project_style.font.name = "Microsoft YaHei"
    project_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    project_style.font.size = Pt(11.0)
    project_style.font.bold = True
    project_style.font.color.rgb = RGBColor(23, 49, 82)

    meta_style = styles["MetaLine"]
    meta_style.base_style = normal
    meta_style.font.name = "Microsoft YaHei"
    meta_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    meta_style.font.size = Pt(9.8)
    meta_style.font.color.rgb = MUTED

    p = doc.add_paragraph()
    format_paragraph(p, space_after=2, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(NAME)
    set_run_font(r, size=18.5, bold=True, color=ACCENT)

    p = doc.add_paragraph(style="MetaLine")
    format_paragraph(p, space_after=7, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(META)
    set_run_font(r, size=9.8, color=MUTED)

    def add_section(title):
        p = doc.add_paragraph(style="ResumeSection")
        format_paragraph(p, space_before=5, space_after=3, line=1.0)
        r = p.add_run(title)
        set_run_font(r, size=12.2, bold=True, color=ACCENT)
        add_section_border(p)

    def add_body(text, space_after=2):
        p = doc.add_paragraph(style="Normal")
        format_paragraph(p, space_after=space_after, line=1.12)
        r = p.add_run(text)
        set_run_font(r, size=10.3, color=TEXT)

    add_section("个人介绍")
    add_body(INTRO, space_after=3)

    add_section("基本信息")
    table = doc.add_table(rows=2, cols=3)
    table.autofit = False
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.width = Cm(5.3)
            key, value = INFO_ROWS[i][j]
            cell.text = ""
            p = cell.paragraphs[0]
            format_paragraph(p, line=1.0)
            r1 = p.add_run(f"{key}：")
            set_run_font(r1, size=10.0, bold=True, color=ACCENT)
            r2 = p.add_run(value)
            set_run_font(r2, size=10.0, color=TEXT)
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F7FAFD")
            tc_pr.append(shd)
    doc.add_paragraph("")

    add_section("教育背景")
    p = doc.add_paragraph(style="Normal")
    format_paragraph(p, space_after=1, line=1.0)
    r = p.add_run(EDU)
    set_run_font(r, size=10.6, bold=True, color=RGBColor(23, 49, 82))
    add_body(EDU_DETAIL, space_after=3)

    add_section("掌握技能")
    for item in SKILLS:
        add_body(item, space_after=1)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(3)

    add_section("项目经历")
    for title, desc in PROJECTS:
        p = doc.add_paragraph(style="ProjectTitle")
        format_paragraph(p, space_before=1, space_after=1, line=1.0)
        r = p.add_run(title)
        set_run_font(r, size=11.0, bold=True, color=RGBColor(23, 49, 82))
        add_body(desc, space_after=2)

    add_section("竞赛经历")
    for item in COMPETITIONS:
        p = doc.add_paragraph(style="List Bullet")
        format_paragraph(p, space_after=1, line=1.0)
        r = p.add_run(item)
        set_run_font(r, size=10.2, color=TEXT)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("简历优化版｜实践项目版")
    set_run_font(fr, size=8.8, color=MUTED)

    doc.save(str(DOCX_PATH))


def build_pdf():
    font_path = r"C:\Windows\Fonts\SourceHanSansCN-Normal.ttf"
    pdfmetrics.registerFont(TTFont("ResumeCN", font_path))

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="TitleCN",
            fontName="ResumeCN",
            fontSize=19,
            leading=22,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#1C497B"),
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="MetaCN",
            fontName="ResumeCN",
            fontSize=9.5,
            leading=12,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#5A6676"),
            spaceAfter=7,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SectionCN",
            fontName="ResumeCN",
            fontSize=12.2,
            leading=15,
            textColor=colors.HexColor("#1C497B"),
            spaceBefore=4,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyCN",
            fontName="ResumeCN",
            fontSize=9.4,
            leading=13.1,
            textColor=colors.HexColor("#1A2330"),
            spaceAfter=2,
            alignment=TA_LEFT,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ProjectCN",
            fontName="ResumeCN",
            fontSize=10.3,
            leading=12.7,
            textColor=colors.HexColor("#173152"),
            spaceBefore=1,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BulletCN",
            fontName="ResumeCN",
            fontSize=9.4,
            leading=12.4,
            textColor=colors.HexColor("#1A2330"),
            leftIndent=10,
            bulletIndent=0,
            spaceAfter=1,
        )
    )

    story = [
        Paragraph(NAME, styles["TitleCN"]),
        Paragraph(META, styles["MetaCN"]),
    ]

    def section(title):
        story.append(Paragraph(title, styles["SectionCN"]))
        story.append(
            HRFlowable(
                width="100%",
                thickness=0.7,
                color=colors.HexColor("#B7CBE3"),
                spaceBefore=1,
                spaceAfter=4,
            )
        )

    section("个人介绍")
    story.append(Paragraph(INTRO, styles["BodyCN"]))

    section("基本信息")
    info_data = []
    for row in INFO_ROWS:
        info_row = []
        for key, value in row:
            info_row.append(Paragraph(f"<b>{key}：</b>{value}", styles["BodyCN"]))
        info_data.append(info_row)
    table = Table(info_data, colWidths=[55 * mm, 55 * mm, 55 * mm])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7FAFD")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#D9E5F2")),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D9E5F2")),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 4))

    section("教育背景")
    story.append(Paragraph(f"<b>{EDU}</b>", styles["BodyCN"]))
    story.append(Paragraph(EDU_DETAIL, styles["BodyCN"]))

    section("掌握技能")
    for item in SKILLS:
        story.append(Paragraph(item, styles["BodyCN"]))

    section("项目经历")
    for title, desc in PROJECTS:
        story.append(
            KeepTogether(
                [
                    Paragraph(f"<b>{title}</b>", styles["ProjectCN"]),
                    Paragraph(desc, styles["BodyCN"]),
                ]
            )
        )

    section("竞赛经历")
    for item in COMPETITIONS:
        story.append(Paragraph(item, styles["BulletCN"], bulletText="•"))

    def add_footer(canvas, _doc):
        canvas.saveState()
        canvas.setFont("ResumeCN", 8.5)
        canvas.setFillColor(colors.HexColor("#5A6676"))
        canvas.drawCentredString(A4[0] / 2.0, 8 * mm, "简历优化版｜实践项目版")
        canvas.restoreState()

    pdf = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=14 * mm,
        rightMargin=14 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
    )
    pdf.build(story, onFirstPage=add_footer, onLaterPages=add_footer)


if __name__ == "__main__":
    ROOT.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
