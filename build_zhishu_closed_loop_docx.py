from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUTPUT = Path("知枢_企业知识资产管理与智能检索平台_闭环落地版_修订稿.docx")
ER_DIAGRAM_DIR = Path("tmp/er_diagrams")
ACCENT = RGBColor(31, 78, 121)
ACCENT_LIGHT = "DCE6F1"
ACCENT_BAND = "EDF3F9"
WARNING_FILL = "FBE5D6"
TEXT_DARK = RGBColor(32, 32, 32)
ER_BG = (247, 250, 252)
ER_HEADER = (79, 129, 189)
ER_HEADER_TEXT = (255, 255, 255)
ER_BORDER = (47, 93, 135)
ER_TEXT = (32, 32, 32)
ER_LINE = (84, 97, 112)
ER_PANEL = (255, 255, 255)
ER_LABEL_FILL = (237, 243, 249)
ER_NOTE = (96, 96, 96)


@dataclass
class TableSpec:
    headers: Sequence[str]
    rows: Sequence[Sequence[str]]
    widths: Sequence[float] | None = None


@dataclass(frozen=True)
class EntitySpec:
    key: str
    title: str
    center: tuple[int, int]
    attributes: dict[str, Sequence[str]]
    width: int = 170
    height: int = 72


@dataclass(frozen=True)
class RelationshipSpec:
    key: str
    title: str
    center: tuple[int, int]
    width: int = 110
    height: int = 72


@dataclass(frozen=True)
class LinkSpec:
    entity_key: str
    relationship_key: str
    entity_card: str
    relationship_card: str


def load_font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = (
        Path(r"C:\Windows\Fonts\msyhbd.ttc") if bold else Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf") if bold else Path(r"C:\Windows\Fonts\simsun.ttc"),
    )
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def measure_text(draw: ImageDraw.ImageDraw, text: str, font) -> tuple[int, int]:
    left, top, right, bottom = draw.textbbox((0, 0), text, font=font)
    return right - left, bottom - top


def draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill) -> None:
    x1, y1, x2, y2 = box
    text_w, text_h = measure_text(draw, text, font)
    draw.text(
        (x1 + (x2 - x1 - text_w) / 2, y1 + (y2 - y1 - text_h) / 2),
        text,
        font=font,
        fill=fill,
    )


def rect_bounds(center: tuple[int, int], width: int, height: int) -> tuple[int, int, int, int]:
    cx, cy = center
    half_w = width // 2
    half_h = height // 2
    return (cx - half_w, cy - half_h, cx + half_w, cy + half_h)


def ellipse_bounds(center: tuple[int, int], width: int, height: int) -> tuple[int, int, int, int]:
    return rect_bounds(center, width, height)


def rectangle_border_point(center: tuple[int, int], width: int, height: int, target: tuple[int, int]) -> tuple[int, int]:
    cx, cy = center
    tx, ty = target
    dx = tx - cx
    dy = ty - cy
    if dx == 0 and dy == 0:
        return center
    scale = max(abs(dx) / (width / 2), abs(dy) / (height / 2))
    return int(cx + dx / scale), int(cy + dy / scale)


def diamond_border_point(center: tuple[int, int], width: int, height: int, target: tuple[int, int]) -> tuple[int, int]:
    cx, cy = center
    tx, ty = target
    dx = tx - cx
    dy = ty - cy
    if dx == 0 and dy == 0:
        return center
    scale = abs(dx) / (width / 2) + abs(dy) / (height / 2)
    return int(cx + dx / scale), int(cy + dy / scale)


def ellipse_border_point(center: tuple[int, int], width: int, height: int, target: tuple[int, int]) -> tuple[int, int]:
    cx, cy = center
    tx, ty = target
    dx = tx - cx
    dy = ty - cy
    if dx == 0 and dy == 0:
        return center
    scale = ((dx * dx) / ((width / 2) ** 2) + (dy * dy) / ((height / 2) ** 2)) ** 0.5
    return int(cx + dx / scale), int(cy + dy / scale)


def draw_entity(draw: ImageDraw.ImageDraw, spec: EntitySpec, *, font) -> tuple[int, int, int, int]:
    bounds = rect_bounds(spec.center, spec.width, spec.height)
    draw.rectangle(bounds, outline=(20, 20, 20), width=3, fill=(255, 255, 255))
    draw_centered_text(draw, bounds, spec.title, font, (20, 20, 20))
    return bounds


def draw_relationship(draw: ImageDraw.ImageDraw, spec: RelationshipSpec, *, font) -> tuple[tuple[int, int], ...]:
    cx, cy = spec.center
    half_w = spec.width // 2
    half_h = spec.height // 2
    points = ((cx, cy - half_h), (cx + half_w, cy), (cx, cy + half_h), (cx - half_w, cy))
    draw.polygon(points, outline=(20, 20, 20), fill=(255, 255, 255), width=3)
    draw_centered_text(draw, rect_bounds(spec.center, spec.width - 10, spec.height - 10), spec.title, font, (20, 20, 20))
    return points


def attribute_centers(spec: EntitySpec) -> list[tuple[str, tuple[int, int]]]:
    centers: list[tuple[str, tuple[int, int]]] = []
    side_offsets = {
        "top": (-1, -118),
        "bottom": (-1, 118),
        "left": (-165, -1),
        "right": (165, -1),
    }
    side_gap = {"top": 118, "bottom": 118, "left": 78, "right": 78}

    for side, labels in spec.attributes.items():
        if not labels:
            continue
        count = len(labels)
        for idx, label in enumerate(labels):
            if side in ("top", "bottom"):
                gap = side_gap[side]
                start = spec.center[0] - ((count - 1) * gap) / 2
                x = int(start + idx * gap)
                y = spec.center[1] + side_offsets[side][1]
            else:
                gap = side_gap[side]
                start = spec.center[1] - ((count - 1) * gap) / 2
                x = spec.center[0] + side_offsets[side][0]
                y = int(start + idx * gap)
            centers.append((label, (x, y)))
    return centers


def draw_attribute(draw: ImageDraw.ImageDraw, center: tuple[int, int], label: str, *, font) -> tuple[int, int, int, int]:
    text_w, text_h = measure_text(draw, label, font)
    width = max(98, text_w + 34)
    height = max(42, text_h + 20)
    bounds = ellipse_bounds(center, width, height)
    draw.ellipse(bounds, outline=(20, 20, 20), width=2, fill=(255, 255, 255))
    draw_centered_text(draw, bounds, label, font, (20, 20, 20))
    return bounds


def draw_entity_attributes(draw: ImageDraw.ImageDraw, spec: EntitySpec, *, entity_font, attr_font) -> None:
    entity_bounds = rect_bounds(spec.center, spec.width, spec.height)
    for label, center in attribute_centers(spec):
        attr_bounds = draw_attribute(draw, center, label, font=attr_font)
        start = rectangle_border_point(spec.center, spec.width, spec.height, center)
        end = ellipse_border_point(center, attr_bounds[2] - attr_bounds[0], attr_bounds[3] - attr_bounds[1], spec.center)
        draw.line((start, end), fill=(20, 20, 20), width=2)


def draw_cardinality(draw: ImageDraw.ImageDraw, position: tuple[int, int], text: str, font) -> None:
    if not text:
        return
    text_w, text_h = measure_text(draw, text, font)
    draw.rectangle((position[0] - 4, position[1] - 2, position[0] + text_w + 4, position[1] + text_h + 2), fill=(255, 255, 255))
    draw.text(position, text, font=font, fill=(20, 20, 20))


def draw_link(
    draw: ImageDraw.ImageDraw,
    entity_map: dict[str, EntitySpec],
    relationship_map: dict[str, RelationshipSpec],
    link: LinkSpec,
    *,
    card_font,
) -> None:
    entity = entity_map[link.entity_key]
    relationship = relationship_map[link.relationship_key]
    start = rectangle_border_point(entity.center, entity.width, entity.height, relationship.center)
    end = diamond_border_point(relationship.center, relationship.width, relationship.height, entity.center)
    draw.line((start, end), fill=(20, 20, 20), width=3)

    dx = end[0] - start[0]
    dy = end[1] - start[1]
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    nx = -dy / length
    ny = dx / length
    start_card_pos = (int(start[0] + dx * 0.28 + nx * 12), int(start[1] + dy * 0.28 + ny * 12))
    draw_cardinality(draw, start_card_pos, link.entity_card, card_font)


def add_diagram_header(draw: ImageDraw.ImageDraw, title: str, subtitle: str) -> None:
    title_font = load_font(42, bold=True)
    subtitle_font = load_font(22)
    draw.text((86, 48), title, font=title_font, fill=ER_BORDER)
    draw.text((88, 110), subtitle, font=subtitle_font, fill=ER_NOTE)
    draw.line((86, 148, 1780, 148), fill=(210, 220, 230), width=3)


def add_diagram_footer(draw: ImageDraw.ImageDraw, text: str) -> None:
    footer_font = load_font(20)
    draw.line((86, 1020, 1780, 1020), fill=(210, 220, 230), width=3)
    draw.text((88, 1036), text, font=footer_font, fill=ER_NOTE)


def render_er_diagram(
    path: Path,
    *,
    title: str,
    subtitle: str,
    footer: str,
    entities: Sequence[EntitySpec],
    relationships: Sequence[RelationshipSpec],
    links: Sequence[LinkSpec],
) -> None:
    image = Image.new("RGB", (1860, 1100), (255, 255, 255))
    draw = ImageDraw.Draw(image)
    entity_font = load_font(22, bold=True)
    relation_font = load_font(20, bold=True)
    attr_font = load_font(18)
    card_font = load_font(20, bold=True)

    add_diagram_header(draw, title, subtitle)

    entity_map = {entity.key: entity for entity in entities}
    relationship_map = {relation.key: relation for relation in relationships}

    for entity in entities:
        draw_entity(draw, entity, font=entity_font)
    for relationship in relationships:
        draw_relationship(draw, relationship, font=relation_font)
    for link in links:
        draw_link(draw, entity_map, relationship_map, link, card_font=card_font)
    for entity in entities:
        draw_entity_attributes(draw, entity, entity_font=entity_font, attr_font=attr_font)

    add_diagram_footer(draw, footer)
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path)


def build_er_diagrams() -> dict[str, Path]:
    diagrams = {
        "er1": ER_DIAGRAM_DIR / "er-1-core-knowledge.png",
        "er2": ER_DIAGRAM_DIR / "er-2-qa-traceability.png",
    }

    render_er_diagram(
        diagrams["er1"],
        title="ER-1 基础知识管理域",
        subtitle="概念层只保留真正独立的核心实体：角色、用户、分类、文档、标签、文档版本，边上标注 1 或 N。",
        footer="图示重点：先讲清谁在维护文档、文档属于哪类、文档如何打标签以及如何形成版本。",
        entities=[
            EntitySpec("roles", "角色", (360, 280), {"top": ["角色编号"], "left": ["角色名称"], "right": ["角色说明"]}),
            EntitySpec("users", "用户", (360, 780), {"left": ["用户编号", "部门"], "bottom": ["用户名"]}),
            EntitySpec("categories", "分类", (940, 280), {"top": ["分类编号"], "left": ["分类名称"], "right": ["分类说明"]}),
            EntitySpec("documents", "文档", (940, 560), {"left": ["文档编号"], "bottom": ["标题"], "right": ["状态"]}),
            EntitySpec("tags", "标签", (1500, 280), {"top": ["标签编号"], "left": ["标签名称"], "right": ["标签说明"]}),
            EntitySpec("versions", "文档版本", (1500, 780), {"left": ["版本编号"], "bottom": ["版本号"], "right": ["变更说明"]}, width=190),
        ],
        relationships=[
            RelationshipSpec("owns", "拥有", (360, 530)),
            RelationshipSpec("creates", "创建", (650, 670)),
            RelationshipSpec("belongs", "归属", (940, 420)),
            RelationshipSpec("binds", "绑定", (1220, 420)),
            RelationshipSpec("derives", "形成版本", (1220, 670), width=128),
        ],
        links=[
            LinkSpec("roles", "owns", "1", "N"),
            LinkSpec("users", "owns", "N", "1"),
            LinkSpec("users", "creates", "1", "N"),
            LinkSpec("documents", "creates", "N", "1"),
            LinkSpec("documents", "belongs", "N", "1"),
            LinkSpec("categories", "belongs", "1", "N"),
            LinkSpec("documents", "binds", "N", "N"),
            LinkSpec("tags", "binds", "N", "N"),
            LinkSpec("documents", "derives", "1", "N"),
            LinkSpec("versions", "derives", "N", "1"),
        ],
    )

    render_er_diagram(
        diagrams["er2"],
        title="ER-2 问答与版本追溯域",
        subtitle="概念层只保留问答闭环里最关键的对象：用户、问题、回答、引用证据、文档、文档版本，边上标注 1 或 N。",
        footer="图示重点：回答不是黑盒文本，它可以通过引用证据回溯到具体文档版本。",
        entities=[
            EntitySpec("users", "用户", (260, 320), {"left": ["用户编号"], "bottom": ["用户名"]}),
            EntitySpec("questions", "问题", (720, 320), {"top": ["问题编号", "问题内容"], "right": ["状态"]}),
            EntitySpec("answers", "回答", (720, 620), {"left": ["回答编号"], "right": ["模型", "回答时间"]}),
            EntitySpec("citations", "引用证据", (720, 860), {"left": ["引用编号"], "bottom": ["证据顺序"]}, width=190),
            EntitySpec("documents", "文档", (1430, 340), {"top": ["文档编号"], "right": ["标题"]}),
            EntitySpec("versions", "文档版本", (1430, 780), {"right": ["版本编号"], "bottom": ["版本号"]}, width=190),
        ],
        relationships=[
            RelationshipSpec("asks", "提出", (470, 320)),
            RelationshipSpec("generates", "生成回答", (720, 470), width=128),
            RelationshipSpec("cites", "引用", (720, 740)),
            RelationshipSpec("version_from", "形成版本", (1430, 560), width=128),
            RelationshipSpec("locate_version", "定位版本", (1080, 860), width=128),
        ],
        links=[
            LinkSpec("users", "asks", "1", "N"),
            LinkSpec("questions", "asks", "N", "1"),
            LinkSpec("questions", "generates", "1", "N"),
            LinkSpec("answers", "generates", "N", "1"),
            LinkSpec("answers", "cites", "1", "N"),
            LinkSpec("citations", "cites", "N", "1"),
            LinkSpec("documents", "version_from", "1", "N"),
            LinkSpec("versions", "version_from", "N", "1"),
            LinkSpec("citations", "locate_version", "N", "1"),
            LinkSpec("versions", "locate_version", "1", "N"),
        ],
    )

    return diagrams


def set_page_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)


def set_run_font(run, size: float = 10.5, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def style_paragraph(paragraph, *, space_before=0, space_after=8, line_spacing=1.35, align=None) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = line_spacing
    if align is not None:
        paragraph.alignment = align


def add_paragraph(
    doc: Document,
    text: str,
    *,
    size: float = 10.5,
    bold: bool = False,
    color: RGBColor | None = None,
    style: str | None = None,
    align=None,
    space_before=0,
    space_after=8,
    line_spacing=1.35,
) -> None:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    style_paragraph(
        paragraph,
        space_before=space_before,
        space_after=space_after,
        line_spacing=line_spacing,
        align=align,
    )


def add_bullet(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)
        style_paragraph(p, space_after=4)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def format_cell(cell, *, header=False, fill: str | None = None, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    if fill:
        shade_cell(cell, fill)
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        style_paragraph(paragraph, space_after=2, line_spacing=1.2)
        for run in paragraph.runs:
            set_run_font(run, size=10, bold=header, color=RGBColor(255, 255, 255) if header else None)


def add_table(doc: Document, spec: TableSpec, *, title: str | None = None, notes: Sequence[str] | None = None) -> None:
    if title:
        add_paragraph(doc, title, size=11.5, bold=True, color=ACCENT, space_before=6, space_after=6)

    table = doc.add_table(rows=1, cols=len(spec.headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    if spec.widths:
        for idx, width in enumerate(spec.widths):
            table.columns[idx].width = Inches(width)

    header_row = table.rows[0]
    for idx, text in enumerate(spec.headers):
        cell = header_row.cells[idx]
        cell.text = text
        format_cell(cell, header=True, fill="4F81BD", align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_index, row_values in enumerate(spec.rows):
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.text = value
            fill = ACCENT_BAND if row_index % 2 == 0 else None
            format_cell(cell, fill=fill, align=WD_ALIGN_PARAGRAPH.LEFT)

    if notes:
        for note in notes:
            add_paragraph(doc, note, size=9.5, color=RGBColor(96, 96, 96), space_before=3, space_after=3, line_spacing=1.2)

    doc.add_paragraph()


def add_section_heading(doc: Document, title: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    size = 15 if level == 1 else 12.5
    set_run_font(run, size=size, bold=True, color=ACCENT)
    style_paragraph(paragraph, space_before=10 if level == 1 else 6, space_after=6, line_spacing=1.15)


def add_callout(doc: Document, title: str, body: str, fill: str = WARNING_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    set_run_font(r1, size=11, bold=True, color=ACCENT)
    style_paragraph(p1, space_after=4, line_spacing=1.1)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=TEXT_DARK)
    style_paragraph(p2, space_after=0, line_spacing=1.25)
    doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, caption: str, note: str) -> None:
    doc.add_picture(str(image_path), width=Cm(16.5))
    picture_paragraph = doc.paragraphs[-1]
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(picture_paragraph, space_before=4, space_after=4, line_spacing=1.0)
    add_paragraph(
        doc,
        caption,
        size=10,
        bold=True,
        color=ACCENT,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0,
        space_after=2,
        line_spacing=1.1,
    )
    add_paragraph(
        doc,
        note,
        size=9.5,
        color=RGBColor(96, 96, 96),
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=0,
        space_after=8,
        line_spacing=1.25,
    )


def configure_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    para = footer.paragraphs[0]
    para.text = ""
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run("知枢闭环落地版修订稿")
    set_run_font(run, size=9, color=RGBColor(100, 100, 100))


def add_cover(doc: Document) -> None:
    add_paragraph(doc, "知枢", size=24, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_paragraph(
        doc,
        "企业知识资产管理与智能检索平台",
        size=18,
        bold=True,
        color=TEXT_DARK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_paragraph(
        doc,
        "闭环落地版执行手册（修订稿）",
        size=15,
        bold=True,
        color=ACCENT,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=16,
    )

    add_callout(
        doc,
        "修订定位",
        "本稿不是对原 PDF 的局部纠错，而是将数据模型、ER 关系、RAG 流程、Agent 留痕、接口与部署方案重写为可直接落地实施的闭环设计。",
    )

    spec = TableSpec(
        headers=["项目项", "内容"],
        rows=[
            ["适用对象", "课程设计、毕业设计、中小型企业知识库原型项目、内部答辩演示稿"],
            ["建议技术栈", "React + TypeScript + Rust(Axum) + MySQL 8.0 + Qdrant + Redis + MinIO + OpenAI-Compatible API"],
            ["本稿核心目标", "可开发、可答辩、可运维、可扩展；同时避免原版中数据模型不闭环的问题"],
            ["修订重点", "多文档引用、版本一致性、向量分段落库、FAQ 来源闭环、Agent 审计可追踪、接口与页面对齐"],
        ],
        widths=[1.8, 4.9],
    )
    add_table(doc, spec)

    doc.add_page_break()


def add_contents(doc: Document) -> None:
    add_section_heading(doc, "目录", level=1)
    for item in [
        "1. 修订原则与总体结论",
        "2. 业务边界、角色与总体架构",
        "3. 数据建模总原则",
        "4. ER 逻辑重构方案",
        "5. 关系模型汇总与关键表详细设计",
        "6. 版本、分段、向量与 Agent 执行闭环",
        "7. 核心业务流程设计",
        "8. 接口与页面闭环映射",
        "9. 安全、部署、测试与实施计划",
        "10. 与原稿相比的关键修订清单",
    ]:
        add_paragraph(doc, item, size=11, space_after=4)
    doc.add_page_break()


def add_section_1(doc: Document) -> None:
    add_section_heading(doc, "1. 修订原则与总体结论", level=1)
    add_paragraph(
        doc,
        "原稿的业务方向正确，但在数据模型与实现闭环上存在三类核心缺口：一是回答记录只能引用单篇文档，无法支撑标准 RAG 场景；二是文档主表与版本表缺乏明确的一致性规则；三是 Agent 处理记录被强绑定到单篇文档，无法完整覆盖问答、检索与审计链路。",
    )
    add_bullet(
        doc,
        [
            "把“文档、版本、分段、向量、问答、引用、Agent”拆成边界明确的七类对象。",
            "把“一个回答引用多篇文档/多个片段”作为默认设计，而不是例外场景。",
            "把“文档当前快照”和“文档历史版本”同时保留，但通过事务规则保证一致性。",
            "把 MinIO、Qdrant、MySQL 的职责分开：文件进对象存储、结构化进 MySQL、向量进 Qdrant。",
            "把 Agent 留痕改成“按执行运行记录 run 追踪”，并允许挂接文档、版本、问题、答案等多类上下文。",
        ],
    )

    add_callout(
        doc,
        "总体结论",
        "建议最终答辩稿采用“4 张 ER 子图 + 1 套关系模型总表 + 5 张关键表详细字段表”的方式呈现。这样既能保证视觉清晰，也能保证老师追问时能落到具体实现层面。",
        fill=ACCENT_LIGHT,
    )


def add_section_2(doc: Document) -> None:
    add_section_heading(doc, "2. 业务边界、角色与总体架构", level=1)
    role_spec = TableSpec(
        headers=["角色", "职责", "关键权限"],
        rows=[
            ["系统管理员", "维护平台级配置、用户与角色、分类、系统参数", "用户管理、角色管理、分类管理、全部文档查看、审计查看"],
            ["知识管理员", "录入文档、维护标签、版本、FAQ，触发知识增强处理", "文档 CRUD、标签管理、FAQ 管理、触发 Agent"],
            ["普通员工", "搜索知识、阅读文档、收藏、提问问答", "查看已发布文档、收藏、提问、查看问答历史"],
        ],
        widths=[1.3, 2.4, 2.7],
    )
    add_table(doc, role_spec, title="2.1 角色定义")

    add_table(
        doc,
        TableSpec(
            headers=["纳入范围", "说明"],
            rows=[
                ["文档录入与发布", "支持正文录入、上传原文、分类、标签、发布与归档"],
                ["版本管理", "每次内容变更写入版本快照并可追溯"],
                ["RAG 问答", "问题入库、结构化检索、向量召回、答案生成、引用回显"],
                ["知识增强", "摘要、标签推荐、FAQ 生成、相关推荐、审计日志"],
                ["统计看板", "文档总量、阅读、收藏、问答趋势、Agent 执行情况"],
            ],
            widths=[2.0, 4.3],
        ),
        title="2.2 范围边界",
        notes=["本版仍不纳入复杂审批流、LDAP/SSO、多租户隔离、OCR 解析、生产级审计合规。"],
    )

    add_table(
        doc,
        TableSpec(
            headers=["层次", "核心职责", "落地组件"],
            rows=[
                ["接入层", "前端页面、鉴权、统一错误提示、交互编排", "React + TypeScript + Ant Design + Axios"],
                ["应用层", "业务编排、权限控制、统一响应模型", "Rust + Axum + Tokio + Serde"],
                ["数据层", "结构化数据、缓存、文件、向量", "MySQL 8.0 + Redis 7 + MinIO + Qdrant"],
                ["智能层", "摘要、标签、FAQ、检索、问答、审计", "OpenAI-Compatible API + 内部 Agent 编排"],
            ],
            widths=[1.2, 2.6, 2.5],
        ),
        title="2.3 总体架构闭环",
    )


def add_section_3(doc: Document) -> None:
    add_section_heading(doc, "3. 数据建模总原则", level=1)
    add_bullet(
        doc,
        [
            "主数据与行为数据分离：文档、用户、标签、分类属于主数据；阅读、收藏、问答、Agent 运行属于行为数据。",
            "当前快照与历史快照分离：documents 存当前可读快照，document_versions 存不可变历史快照。",
            "问答与引用分离：answers 只描述答案本体，answer_citations 单独保存证据片段与排序。",
            "对象存储与结构化分离：MinIO 只保管原文件对象，MySQL 只保管元数据与业务关系。",
            "向量与结构化双落点：Qdrant 保存向量与 payload，MySQL 保存 segment 元数据、排序、版本归属和失效状态。",
        ],
    )
    add_callout(
        doc,
        "约束原则",
        "凡是用户界面需要稳定展示、筛选、排序、审计的内容，必须在 MySQL 中可定位；凡是大模型或向量检索的运行证据，必须能回溯到具体文档版本与片段。",
        fill=ACCENT_LIGHT,
    )


def add_section_4(doc: Document) -> None:
    add_section_heading(doc, "4. ER 逻辑重构方案", level=1)
    add_paragraph(
        doc,
        "本稿将概念 ER 图压缩为 2 张，只保留真正具有独立业务意义的核心实体。像阅读时间、收藏时间、文件来源、Agent 运行状态这类更适合作为属性、外键或实现层表字段的内容，不再在概念图里单独占一个实体位置。",
    )

    er_maps = TableSpec(
        headers=["子图", "覆盖实体", "核心关系", "拆分理由"],
        rows=[
            ["ER-1 基础知识管理域", "角色 / 用户 / 分类 / 文档 / 标签 / 文档版本", "角色-用户 1:N；用户-文档 1:N；分类-文档 1:N；文档-标签 N:N；文档-版本 1:N", "保留知识管理主干，去掉阅读记录、收藏记录等实现层行为实体"],
            ["ER-2 问答与版本追溯域", "用户 / 问题 / 回答 / 引用证据 / 文档 / 文档版本", "用户-问题 1:N；问题-回答 1:N；回答-引用证据 1:N；文档-版本 1:N；引用证据-文档版本 N:1", "保留问答解释链，去掉片段、文件、Agent 运行等实现层细节实体"],
        ],
        widths=[1.5, 2.3, 2.3, 1.9],
    )
    add_table(doc, er_maps, title="4.1 ER 核心子图方案")
    add_callout(
        doc,
        "绘图原则",
        "概念 ER 图只回答“系统里有哪些核心对象、它们之间是什么关系”。实现层的日志、对象存储、向量分段等内容可以在关系模式或实现设计里体现，但不再占用概念实体位置。",
        fill=ACCENT_LIGHT,
    )

    diagrams = build_er_diagrams()
    add_figure(
        doc,
        diagrams["er1"],
        "图 4-1  基础知识管理域 ER 图",
        "讲解顺序建议从角色、用户、分类、文档讲起，再补标签和文档版本，先把知识库主干讲清楚。",
    )
    add_figure(
        doc,
        diagrams["er2"],
        "图 4-2  问答与版本追溯域 ER 图",
        "该图强调回答不是直接挂在文档上，而是通过引用证据回溯到具体文档版本，从而体现答案可解释性。",
    )

    relationship_spec = TableSpec(
        headers=["联系名称", "参与实体", "基数", "落地方式"],
        rows=[
            ["拥有", "角色 - 用户", "1:N", "用户.角色编号 -> 角色.角色编号"],
            ["创建", "用户 - 文档", "1:N", "文档.创建者编号 -> 用户.用户编号"],
            ["归属", "分类 - 文档", "1:N", "文档.分类编号 -> 分类.分类编号"],
            ["绑定", "文档 - 标签", "N:N", "通过 文档标签(文档编号, 标签编号) 转换"],
            ["形成版本", "文档 - 文档版本", "1:N", "文档版本.文档编号 -> 文档.文档编号"],
            ["提出", "用户 - 问题", "1:N", "问题.用户编号 -> 用户.用户编号"],
            ["生成回答", "问题 - 回答", "1:N", "回答.问题编号 -> 问题.问题编号"],
            ["引用证据", "回答 - 引用证据", "1:N", "引用证据.回答编号 -> 回答.回答编号"],
            ["定位版本", "引用证据 - 文档版本", "N:1", "引用证据.版本编号 -> 文档版本.版本编号"],
        ],
        widths=[1.1, 1.8, 0.8, 3.0],
    )
    add_table(doc, relationship_spec, title="4.2 联系说明总表")


def add_section_5(doc: Document) -> None:
    add_section_heading(doc, "5. 关系模型汇总与关键表详细设计", level=1)
    add_paragraph(
        doc,
        "本节先给出与 ER 图完全对应的中文关系模式，再给出落地到物理表后的关键表设计。转换原则是：`1..1 / 1..N` 联系并入 `N` 端实体；`1..N / 1..N` 联系单独转换为中间关系模式。",
    )
    relation_modes = TableSpec(
        headers=["类别", "中文关系模式", "说明"],
        rows=[
            ["实体", "角色(角色编号 PK，角色名称，角色说明)", "对应 ER-1 中的角色实体"],
            ["实体", "用户(用户编号 PK，角色编号 FK，用户名，部门)", "“拥有”联系并入用户，角色编号作为外键"],
            ["实体", "分类(分类编号 PK，分类名称，分类说明)", "对应 ER-1 中的分类实体"],
            ["实体", "标签(标签编号 PK，标签名称，标签说明)", "对应 ER-1 中的标签实体"],
            ["实体", "文档(文档编号 PK，分类编号 FK，创建者编号 FK，标题，状态)", "“归属”“创建”联系并入文档"],
            ["实体", "文档版本(版本编号 PK，文档编号 FK，版本号，变更说明)", "“形成版本”联系并入文档版本"],
            ["实体", "问题(问题编号 PK，用户编号 FK，问题内容，状态)", "“提出”联系并入问题"],
            ["实体", "回答(回答编号 PK，问题编号 FK，模型，回答时间)", "“生成回答”联系并入回答"],
            ["实体", "引用证据(引用编号 PK，回答编号 FK，版本编号 FK，证据顺序)", "回答通过引用证据定位到文档版本"],
            ["联系", "文档标签(文档编号 PK/FK，标签编号 PK/FK)", "“文档-标签”是多对多联系，需单独转换"],
        ],
        widths=[0.8, 4.3, 1.9],
    )
    add_table(doc, relation_modes, title="5.1 中文关系模式表")

    summary = TableSpec(
        headers=["物理表", "用途", "关键约束"],
        rows=[
            ["roles", "角色主档", "role_name UNIQUE"],
            ["users", "用户主档", "username UNIQUE；email 可唯一"],
            ["categories", "分类主档", "category_name UNIQUE"],
            ["document_files", "原始上传文件元数据", "object_key UNIQUE；sha256 UNIQUE 可选"],
            ["documents", "文档当前快照", "category_id / creator_id FK；status + current_version_no"],
            ["document_versions", "文档历史快照", "document_id + version_no UNIQUE"],
            ["tags", "标签主档", "tag_name UNIQUE"],
            ["document_tags", "文档标签关联", "document_id + tag_id UNIQUE"],
            ["read_records", "阅读行为日志", "按需保留重复记录，用于行为统计"],
            ["favorite_records", "收藏状态表", "user_id + document_id UNIQUE"],
            ["questions", "问题主表", "user_id FK；status 字段标记回答状态"],
            ["answers", "回答主表", "question_id FK；一问多答可保留版本"],
            ["answer_citations", "答案引用证据", "answer_id + cite_order UNIQUE"],
            ["faq_items", "FAQ 条目", "document_id FK；status 可选"],
            ["document_segments", "分段元数据表", "version_id + chunk_order UNIQUE；segment_id 与 Qdrant payload 对齐"],
            ["agent_runs", "Agent 执行运行日志", "agent_type + status + started_at 索引"],
        ],
        widths=[1.8, 2.3, 2.4],
    )
    add_table(doc, summary, title="5.2 关系模型汇总表")

    add_table(
        doc,
        TableSpec(
            headers=["字段", "类型", "约束", "说明"],
            rows=[
                ["document_id", "BIGINT", "PK", "文档主键"],
                ["category_id", "BIGINT", "FK -> categories", "所属分类"],
                ["creator_id", "BIGINT", "FK -> users", "创建者"],
                ["current_version_id", "BIGINT", "FK -> document_versions", "当前可读版本"],
                ["current_version_no", "VARCHAR(20)", "NOT NULL", "当前版本号，如 v1.0.0"],
                ["title", "VARCHAR(200)", "NOT NULL", "文档标题"],
                ["summary", "TEXT", "NULL", "当前展示摘要"],
                ["status", "VARCHAR(20)", "NOT NULL", "draft / published / archived"],
                ["source_file_id", "BIGINT", "FK -> document_files，可空", "当前版本对应原文件"],
                ["published_at", "DATETIME", "NULL", "首次发布或最近发布时刻"],
                ["created_at / updated_at", "DATETIME", "NOT NULL", "创建与更新时间"],
            ],
            widths=[1.5, 1.2, 1.6, 2.4],
        ),
        title="5.3 documents（文档当前快照）",
        notes=["说明：documents 不保存历史正文的唯一来源，而是保存当前可读快照；历史内容以 document_versions 为准。"],
    )

    add_table(
        doc,
        TableSpec(
            headers=["字段", "类型", "约束", "说明"],
            rows=[
                ["version_id", "BIGINT", "PK", "版本快照主键"],
                ["document_id", "BIGINT", "FK -> documents", "所属文档"],
                ["version_no", "VARCHAR(20)", "UNIQUE(document_id, version_no)", "版本号"],
                ["title", "VARCHAR(200)", "NOT NULL", "该版本标题快照"],
                ["content", "LONGTEXT", "NOT NULL", "正文快照"],
                ["summary", "TEXT", "NULL", "该版本摘要快照"],
                ["change_note", "VARCHAR(255)", "NOT NULL", "本次变更说明"],
                ["source_file_id", "BIGINT", "FK -> document_files，可空", "版本来源文件"],
                ["is_published_snapshot", "TINYINT(1)", "NOT NULL", "是否发布快照"],
                ["created_by", "BIGINT", "FK -> users", "创建本版本的操作人"],
                ["created_at", "DATETIME", "NOT NULL", "版本生成时间"],
            ],
            widths=[1.4, 1.1, 1.8, 2.4],
        ),
        title="5.4 document_versions（文档历史版本）",
    )

    add_table(
        doc,
        TableSpec(
            headers=["字段", "类型", "约束", "说明"],
            rows=[
                ["answer_id", "BIGINT", "PK", "回答主键"],
                ["question_id", "BIGINT", "FK -> questions", "来源问题"],
                ["answer_text", "LONGTEXT", "NOT NULL", "最终展示答案"],
                ["confidence_score", "DECIMAL(5,2)", "NULL", "可选评分"],
                ["model_name", "VARCHAR(100)", "NOT NULL", "本次回答所用模型"],
                ["status", "VARCHAR(20)", "NOT NULL", "success / partial / failed"],
                ["latency_ms", "INT", "NULL", "整体响应耗时"],
                ["created_at", "DATETIME", "NOT NULL", "回答生成时间"],
            ],
            widths=[1.4, 1.1, 1.7, 2.5],
        ),
        title="5.5 answers（回答主表）",
    )

    add_table(
        doc,
        TableSpec(
            headers=["字段", "类型", "约束", "说明"],
            rows=[
                ["citation_id", "BIGINT", "PK", "引用主键"],
                ["answer_id", "BIGINT", "FK -> answers", "所属回答"],
                ["document_id", "BIGINT", "FK -> documents", "命中文档"],
                ["version_id", "BIGINT", "FK -> document_versions", "命中版本"],
                ["segment_id", "BIGINT", "FK -> document_segments", "命中片段"],
                ["cite_order", "INT", "UNIQUE(answer_id, cite_order)", "展示顺序"],
                ["score", "DECIMAL(8,4)", "NULL", "召回或重排得分"],
                ["snippet_text", "TEXT", "NOT NULL", "前端可直接展示的证据摘要"],
            ],
            widths=[1.4, 1.4, 2.2, 1.8],
        ),
        title="5.6 answer_citations（多证据引用表）",
        notes=["这是本次修订的关键表，用来解决“一个回答只能引用一篇文档”的原始缺陷。"],
    )

    add_table(
        doc,
        TableSpec(
            headers=["字段", "类型", "约束", "说明"],
            rows=[
                ["segment_id", "BIGINT", "PK", "分段主键"],
                ["version_id", "BIGINT", "FK -> document_versions", "所属版本"],
                ["document_id", "BIGINT", "FK -> documents", "冗余保存，便于查询"],
                ["chunk_order", "INT", "UNIQUE(version_id, chunk_order)", "版本内顺序"],
                ["chunk_text", "LONGTEXT", "NOT NULL", "分段正文"],
                ["token_count", "INT", "NULL", "分词或 token 数"],
                ["qdrant_point_id", "VARCHAR(64)", "UNIQUE", "对应向量点标识"],
                ["is_active", "TINYINT(1)", "NOT NULL", "是否参与当前检索"],
                ["created_at", "DATETIME", "NOT NULL", "分段生成时间"],
            ],
            widths=[1.4, 1.4, 2.2, 1.8],
        ),
        title="5.7 document_segments（分段元数据表）",
    )

    add_table(
        doc,
        TableSpec(
            headers=["字段", "类型", "约束", "说明"],
            rows=[
                ["run_id", "BIGINT", "PK", "运行主键"],
                ["agent_type", "VARCHAR(50)", "NOT NULL", "summary / tag / faq / embedding / retrieval / answer / audit"],
                ["trigger_type", "VARCHAR(30)", "NOT NULL", "manual / document_publish / question_submit / scheduled"],
                ["operator_user_id", "BIGINT", "FK -> users，可空", "触发人"],
                ["document_id", "BIGINT", "FK -> documents，可空", "文档上下文"],
                ["version_id", "BIGINT", "FK -> document_versions，可空", "版本上下文"],
                ["question_id", "BIGINT", "FK -> questions，可空", "问题上下文"],
                ["answer_id", "BIGINT", "FK -> answers，可空", "答案上下文"],
                ["status", "VARCHAR(20)", "NOT NULL", "running / success / failed"],
                ["input_text / output_text", "LONGTEXT", "可空", "审计输入输出"],
                ["meta_json", "JSON", "可空", "候选文档、召回参数、模型配置等结构化元数据"],
                ["started_at / finished_at", "DATETIME", "NOT NULL / 可空", "运行时间窗"],
            ],
            widths=[1.0, 1.3, 1.3, 1.2],
        ),
        title="5.8 agent_runs（执行审计主表）",
        notes=["不再强制要求每条 Agent 记录必须绑定单个 document_id；问答型运行可仅绑定 question_id / answer_id。"],
    )


def add_section_6(doc: Document) -> None:
    add_section_heading(doc, "6. 版本、分段、向量与 Agent 执行闭环", level=1)
    add_table(
        doc,
        TableSpec(
            headers=["场景", "事务规则", "结果"],
            rows=[
                ["新增文档", "写 documents 草稿 -> 同事务写 document_versions v1 -> 生成分段与向量 -> 回写 current_version_id", "文档从第一天起即可追溯"],
                ["编辑草稿", "新建版本快照，不覆盖旧 version；更新 documents 当前快照", "主表展示最新，版本表保留历史"],
                ["发布文档", "校验当前版本完整性 -> 更新 documents.status=published -> 标记版本为发布快照", "前台只读取已发布快照"],
                ["归档文档", "documents.status=archived；segments.is_active=0", "文档保留但退出检索结果"],
            ],
            widths=[1.2, 3.3, 1.9],
        ),
        title="6.1 版本一致性策略",
    )

    add_table(
        doc,
        TableSpec(
            headers=["步骤", "输入", "输出", "落库位置"],
            rows=[
                ["清洗", "文档版本正文", "标准化文本", "临时内存 / meta_json"],
                ["切分", "标准化文本", "若干 chunk", "document_segments"],
                ["向量化", "chunk_text", "embedding 向量", "Qdrant points"],
                ["索引绑定", "segment_id + point_id", "可检索分段", "document_segments.qdrant_point_id + Qdrant payload"],
                ["失效处理", "旧版本 segment_id", "inactive 标记", "document_segments.is_active=0"],
            ],
            widths=[0.8, 1.8, 1.8, 2.8],
        ),
        title="6.2 分段与向量闭环",
        notes=["建议 Qdrant payload 至少冗余：segment_id、document_id、version_id、chunk_order、title、category_id。"],
    )

    add_table(
        doc,
        TableSpec(
            headers=["Agent", "默认上下文", "建议写入 agent_runs 的键"],
            rows=[
                ["摘要 Agent", "document_id + version_id", "agent_type=summary，保留摘要输入与摘要输出"],
                ["标签 Agent", "document_id + version_id", "记录推荐标签列表与最终采纳标签"],
                ["FAQ Agent", "document_id + version_id", "记录生成 FAQ 数量与条目摘要"],
                ["Embedding Agent", "version_id", "记录分段数量、模型名、point_id 列表摘要"],
                ["检索 Agent", "question_id", "记录召回参数、候选 segment_id、重排得分"],
                ["问答 Agent", "question_id + answer_id", "记录答案生成参数、模型、耗时、失败原因"],
                ["审计 Agent", "任意运行 run_id", "对异常输出或敏感输出进行审计复核"],
            ],
            widths=[1.1, 1.7, 3.3],
        ),
        title="6.3 Agent 运行记录规范",
    )


def add_section_7(doc: Document) -> None:
    add_section_heading(doc, "7. 核心业务流程设计", level=1)
    add_bullet(
        doc,
        [
            "文档入库：前端提交元数据与正文/原文件 -> 后端创建 documents 与 document_versions -> 文件写 MinIO -> 触发摘要/标签/FAQ/Embedding -> 写入 tags、faq_items、document_segments、agent_runs。",
            "问答流程：前端提问 -> questions 入库 -> MySQL 标题/摘要/标签预筛 -> Qdrant 召回 -> 重排 -> answers 入库 -> answer_citations 逐条保存引用证据 -> 返回答案与引用。",
            "文档更新：用户编辑正文 -> 新建 document_versions -> 旧 segments 失活 -> 新 segments 生成 -> current_version_id 指向新版本 -> 必要时重新生成摘要/标签。",
            "文档删除：业务上建议使用 archived 归档而不是物理删除；如必须删除，应先清理 Qdrant points、MinIO 对象、再删 MySQL 关系。",
        ],
    )

    add_table(
        doc,
        TableSpec(
            headers=["流程", "前端页面", "后端接口", "关键表"],
            rows=[
                ["新增文档", "DocumentForm", "POST /api/documents", "documents, document_versions, document_files"],
                ["发布文档", "DocumentDetail / Admin", "POST /api/documents/{id}/publish", "documents, document_versions, agent_runs"],
                ["提问问答", "QA Page", "POST /api/qa/ask", "questions, answers, answer_citations, agent_runs"],
                ["查看处理日志", "Agent Records", "GET /api/agent-runs", "agent_runs"],
                ["查看版本", "DocumentDetail", "GET /api/documents/{id}/versions", "document_versions"],
            ],
            widths=[1.1, 1.6, 2.0, 1.8],
        ),
        title="7.1 流程到实现的闭环映射",
    )


def add_section_8(doc: Document) -> None:
    add_section_heading(doc, "8. 接口与页面闭环映射", level=1)
    add_table(
        doc,
        TableSpec(
            headers=["接口", "方法", "作用", "补充说明"],
            rows=[
                ["/api/auth/login", "POST", "登录并返回 JWT", "仅后端保存敏感配置"],
                ["/api/users", "GET/POST/PUT", "用户查询与维护", "管理员权限"],
                ["/api/roles", "GET", "角色列表", "可静态化也可走接口"],
                ["/api/categories", "GET/POST/PUT/DELETE", "分类管理", "建议使用 /api/categories/{id} 形式"],
                ["/api/tags", "GET/POST/PUT/DELETE", "标签管理", "支持关键词搜索"],
                ["/api/documents", "GET/POST", "列表与新增", "支持标题、分类、标签、状态筛选"],
                ["/api/documents/{id}", "GET/PUT", "详情与修改", "返回当前快照与当前标签"],
                ["/api/documents/{id}/versions", "GET", "版本列表", "支持查看版本差异"],
                ["/api/documents/{id}/publish", "POST", "发布文档", "触发发布态校验"],
                ["/api/documents/{id}/archive", "POST", "归档文档", "替代直接 DELETE"],
                ["/api/documents/{id}/favorite", "POST/DELETE", "收藏/取消收藏", "映射 favorite_records"],
                ["/api/documents/{id}/read", "POST", "写入阅读记录", "映射 read_records"],
                ["/api/documents/{id}/summary", "POST", "手动触发摘要", "写 agent_runs"],
                ["/api/documents/{id}/faq", "POST", "手动触发 FAQ", "写 faq_items 与 agent_runs"],
                ["/api/qa/ask", "POST", "提交问题并返回答案", "返回 citations 列表"],
                ["/api/questions/history", "GET", "问答历史", "按当前用户过滤"],
                ["/api/agent-runs", "GET", "查看智能处理记录", "支持 agent_type、status、时间范围筛选"],
            ],
            widths=[2.0, 0.8, 1.8, 2.2],
        ),
        title="8.1 推荐接口清单",
    )

    add_table(
        doc,
        TableSpec(
            headers=["页面", "必须展示的数据", "依赖接口"],
            rows=[
                ["Dashboard", "文档总量、发布量、热门文档、问答趋势、Agent 成功率", "/api/dashboard 或聚合接口"],
                ["DocumentList", "标题、摘要、分类、标签、状态、更新时间", "GET /api/documents"],
                ["DocumentDetail", "正文、版本、标签、FAQ、相关推荐、收藏状态", "GET /api/documents/{id} + versions + favorite"],
                ["DocumentForm", "标题、正文、分类、标签、变更说明、原文件", "POST/PUT /api/documents"],
                ["QA", "问题、答案、引用证据、历史问题", "POST /api/qa/ask + GET /api/questions/history"],
                ["AgentRecords", "运行类型、状态、耗时、输入输出摘要", "GET /api/agent-runs"],
            ],
            widths=[1.3, 3.2, 2.0],
        ),
        title="8.2 页面与接口对齐表",
    )


def add_section_9(doc: Document) -> None:
    add_section_heading(doc, "9. 安全、部署、测试与实施计划", level=1)
    add_bullet(
        doc,
        [
            "所有敏感配置只保留在 backend/.env 或部署平台密钥管理中；前端不得持有任何 LLM、Qdrant、MinIO 密钥。",
            "用户上传文档必须设置大小限制、类型白名单与病毒扫描扩展位；MinIO 对象路径建议按 yyyy/mm/document_id/version_id 组织。",
            "问答链路应设置超时、重试与熔断；检索失败时返回“暂无足够依据”，不要伪造答案。",
            "agent_runs 需记录失败原因、耗时、模型名和触发源，便于调试和答辩演示。",
        ],
    )

    add_table(
        doc,
        TableSpec(
            headers=["测试编号", "测试内容", "预期结果", "关键表或模块"],
            rows=[
                ["TC01", "新增文档并上传原文件", "documents、document_versions、document_files 均成功写入", "文档入库"],
                ["TC02", "修改文档并填写变更说明", "产生新版本且旧版本仍可追溯", "版本管理"],
                ["TC03", "同一用户重复收藏同一文档", "返回幂等成功或提示已收藏，不产生重复记录", "favorite_records 唯一约束"],
                ["TC04", "问答命中多篇文档", "answers 写入 1 条，answer_citations 写入多条", "RAG 闭环"],
                ["TC05", "文档归档后再次搜索", "文档退出检索结果但历史记录保留", "documents / document_segments"],
                ["TC06", "FAQ 生成失败", "agent_runs.status=failed 并记录失败原因", "Agent 审计"],
                ["TC07", "旧版本重新切分后检索", "仅 active segments 参与召回", "document_segments / Qdrant"],
            ],
            widths=[0.8, 2.5, 2.2, 1.7],
        ),
        title="9.1 回归测试矩阵",
    )

    add_table(
        doc,
        TableSpec(
            headers=["阶段", "周期", "交付物"],
            rows=[
                ["阶段 1", "第 1 周", "需求边界、角色、样例资料、关系模型初稿"],
                ["阶段 2", "第 2 周", "ER 图四分图、DDL 脚本、基础 API 定义"],
                ["阶段 3", "第 3-4 周", "前端页面骨架、鉴权、文档 CRUD、版本管理"],
                ["阶段 4", "第 5-6 周", "FAQ / 摘要 / 标签 / 向量检索 / 问答链路"],
                ["阶段 5", "第 7 周", "统计看板、测试报告、答辩讲稿与演示视频"],
            ],
            widths=[1.0, 1.1, 4.8],
        ),
        title="9.2 实施计划",
    )


def add_section_10(doc: Document) -> None:
    add_section_heading(doc, "10. 与原稿相比的关键修订清单", level=1)
    add_table(
        doc,
        TableSpec(
            headers=["原稿问题", "本稿修订", "收益"],
            rows=[
                ["answers 只能引用单篇文档", "新增 answer_citations，支持一答多证据", "符合真实 RAG 问答展示"],
                ["agent_records 强绑定 document_id", "改为 agent_runs，允许 question / answer / version 上下文", "问答链路可完整审计"],
                ["MinIO 只有配置，没有数据模型", "新增 document_files", "文件上传链路闭环"],
                ["分段只在说明里出现，没有数据表", "新增 document_segments", "向量与版本可以追溯"],
                ["documents 与 document_versions 没有事务规则", "定义创建、编辑、发布、归档四类一致性规则", "避免双写不一致"],
                ["FAQ 同时画关系又画来源属性", "统一为 faq_items.document_id 外键表达", "ER 图和关系表一致"],
                ["接口设计未覆盖版本、收藏、问答历史", "补齐接口清单并与页面逐项映射", "实现路径更清晰"],
            ],
            widths=[2.1, 2.3, 1.8],
        ),
        title="10.1 修订摘要",
    )

    add_callout(
        doc,
        "交付建议",
        "如果后续需要继续用于答辩，建议下一步把本稿中的“ER 子图拆分建议”正式绘制成四张图，并补一份 MySQL DDL 附录。这样老师从逻辑层、表结构层到实现层的追问都能闭环。",
        fill=ACCENT_LIGHT,
    )


def build() -> None:
    doc = Document()
    set_page_layout(doc)
    configure_footer(doc)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    add_cover(doc)
    add_contents(doc)
    add_section_1(doc)
    add_section_2(doc)
    add_section_3(doc)
    add_section_4(doc)
    add_section_5(doc)
    add_section_6(doc)
    add_section_7(doc)
    add_section_8(doc)
    add_section_9(doc)
    add_section_10(doc)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
